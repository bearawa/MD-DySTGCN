from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from math2docx import add_math


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "基于气象动态图与分片Transformer的汉江水质参数预测研究.md"
FIGURE_ROOT = ROOT / "论文插图包" / "figures"
SKILL_FIGURE_ROOT = ROOT / "论文撰写" / "skill-generated-figures"
SKILL_FIGURE_OVERRIDES = {
    "图1-1_技术路线.png": SKILL_FIGURE_ROOT / "图1_技术路线_skill.png",
    "图2-1_气象驱动机理.png": SKILL_FIGURE_ROOT / "图2_气象驱动机理_skill.png",
    "图4-1_MD-DySTGCN-CA结构.png": (
        ROOT / "论文插图包" / "figures" / "ch4" / "图4-1_MD-DySTGCN-CA结构.png"
    ),
    "图4-2_动态邻接.png": SKILL_FIGURE_ROOT / "图12_动态邻接_skill.png",
    "图4-3_残差TCN.png": SKILL_FIGURE_ROOT / "图11_残差TCN_skill.png",
    "图4-4_门控融合.png": SKILL_FIGURE_ROOT / "图13_门控融合_skill.png",
}
OUTPUT = ROOT / "论文撰写" / "汉江水质参数预测研究_学年论文_MD-DySTGCN-CA.docx"

MAIN_TITLE = "汉江水质参数短时预测"
SUB_TITLE = "——气象感知动态图与 TCN 交叉注意力融合"
EN_TITLE = (
    "Short-Term Hanjiang River Water-Quality Parameter Prediction "
    "Using Meteorology-Aware Dynamic Graphs and TCN Cross-Attention"
)

FIGURE_MAP = {
    "图1-1": "图1",
    "图2-1": "图2",
    "图3-1": "图3",
    "图3-2": "图4",
    "图3-5": "图5",
    "图3-3": "图6",
    "图3-4": "图7",
    "图3-6": "图8",
    "图3-7": "图9",
    "图4-1": "图10",
    "图4-3": "图11",
    "图4-2": "图12",
    "图4-4": "图13",
    "图5-2b": "图15",
    "图5-2c": "图16",
    "图5-2": "图14",
    "图5-1": "图17",
    "图5-4": "图18",
    "图5-3": "图19",
}

TABLE_MAP = {
    "表3-1": "表1",
    "表3-2": "表2",
    "表3-3": "表3",
    "表3-4": "表4",
    "表5-1": "表5",
    "表5-2": "表6",
    "表5-3": "表7",
    "表5-4": "表8",
}

EQUATION_MAP = {
    "2-1": 1,
    "2-2": 2,
    "2-3": 3,
    "2-4": 4,
    "2-5": 5,
    "2-6": 6,
    "3-1": 7,
    "4-1": 8,
    "4-2": 9,
    "4-3": 10,
    "4-4": 11,
    "4-5": 12,
    "4-6": 13,
    "4-7": 14,
    "4-8": 15,
}

CN_NUMERALS = {
    1: "一",
    2: "二",
    3: "三",
    4: "四",
    5: "五",
    6: "六",
    7: "七",
    8: "八",
    9: "九",
    10: "十",
}

OMIT_CONTAINS = (
    "扫描件",
    "本机复现图",
    "图录见",
    "不以原文",
    "原文图",
    "图5-1（同款）",
    "*_e50/history.json",
)


def set_run_font(run, chinese="宋体", western="Times New Roman", size=Pt(12), bold=None):
    run.font.name = western
    run.font.size = size
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), chinese)
    rfonts.set(qn("w:ascii"), western)
    rfonts.set(qn("w:hAnsi"), western)


def set_cell_shading(cell, fill="E7E6E6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    ppr = paragraph._p.get_or_add_pPr()
    keep = ppr.find(qn("w:keepNext"))
    if value and keep is None:
        ppr.append(OxmlElement("w:keepNext"))
    elif not value and keep is not None:
        ppr.remove(keep)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def clean_markdown(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = text.replace("\\(", "").replace("\\)", "")
    text = text.replace("\\[", "").replace("\\]", "")
    text = normalize_labels(text)
    return latex_to_text(text)


def academic_heading(text: str, level: int) -> str:
    text = clean_markdown(text)
    if level == 1:
        match = re.match(r"第(\d+)章\s*(.*)", text)
        if match:
            number = CN_NUMERALS[int(match.group(1))]
            return f"{number}、{match.group(2)}"
    elif level == 2:
        match = re.match(r"(\d+)\.(\d+)\s*(.*)", text)
        if match:
            number = CN_NUMERALS[int(match.group(2))]
            return f"（{number}）{match.group(3)}"
    elif level == 3:
        match = re.match(r"(\d+)\.(\d+)\.(\d+)\s*(.*)", text)
        if match:
            return f"{int(match.group(3))}．{match.group(4)}"
    return text


def normalize_labels(text: str) -> str:
    text = re.sub(r"图\s*5-1（续）", "图18", text)
    for old in sorted(FIGURE_MAP, key=len, reverse=True):
        text = re.sub(
            re.escape(old).replace("图", r"图\s*"),
            FIGURE_MAP[old],
            text,
        )
    for old in sorted(TABLE_MAP, key=len, reverse=True):
        text = re.sub(
            re.escape(old).replace("表", r"表\s*"),
            TABLE_MAP[old],
            text,
        )
    return text


def latex_to_text(text: str) -> str:
    text = text.replace(r"\left", "").replace(r"\right", "")
    replacements = {
        r"\Delta": "Δ",
        r"\sigma": "σ",
        r"\delta": "δ",
        r"\ell": "ℓ",
        r"\mu": "μ",
        r"\theta": "θ",
        r"\tau": "τ",
        r"\Psi": "Ψ",
        r"\ldots": "…",
        r"\top": "⊤",
        r"\sum": "Σ",
        r"\exp": "exp",
        r"\min": "min",
        r"\max": "max",
        r"\mathcal": "",
        r"\times": "×",
        r"\cdot": "·",
        r"\odot": "⊙",
        r"\in": "∈",
        r"\le": "≤",
        r"\ge": "≥",
        r"\sim": "～",
        r"\rightarrow": "→",
        r"\big": "",
        r"\lvert": "|",
        r"\rvert": "|",
        r"\mathbb{R}": "ℝ",
        r"\mathbf": "",
        r"\mathrm": "",
        r"\text": "",
        r"\quad": "    ",
        r"\,": " ",
        r"\!": "",
        r"\ ": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    while "{{" in text or "}}" in text:
        text = text.replace("{{", "{").replace("}}", "}")
    text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)
    text = re.sub(r"\\hat\s*\{?([A-Za-z])\}?", lambda m: f"{m.group(1)}̂", text)
    text = re.sub(r"\\tilde\{?([A-Za-z])\}?", lambda m: f"{m.group(1)}̃", text)
    text = text.replace(r"\frac12", "(1)/(2)")
    for _ in range(5):
        text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
        text = re.sub(r"\\sqrt\{([^{}]+)\}", r"√(\1)", text)
    text = re.sub(r"\\operatorname\{([^{}]+)\}", r"\1", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("$", "")
    return text


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.first_line_indent = Cm(0.85)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, bold, center, chinese in (
        ("Title", 22, True, True, "黑体"),
        ("Subtitle", 16, False, True, "宋体"),
        ("Heading 1", 16, True, True, "黑体"),
        ("Heading 2", 14, True, False, "宋体"),
        ("Heading 3", 12, True, False, "宋体"),
    ):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
        style.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        )
        style.paragraph_format.first_line_indent = Cm(0)
        if style_name in {"Heading 2", "Heading 3"}:
            style.paragraph_format.left_indent = Cm(0.85)
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for style_name, size, chinese, bold in (
        ("TOC 1", 14, "黑体", True),
        ("TOC 2", 12, "宋体", False),
    ):
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)

    if "Figure Caption" not in styles:
        style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Figure Caption"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.bold = True
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(6)

    if "Reference" not in styles:
        style = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Reference"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.first_line_indent = Cm(0.74)
    style.paragraph_format.line_spacing = 1.25


def configure_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)


def add_centered_text(doc, text, size, bold=False, chinese="宋体", before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, chinese=chinese, size=Pt(size), bold=bold)
    return p


def add_cover(doc: Document):
    add_centered_text(doc, "中南财经政法大学", 18, True, "黑体", before=10)
    add_centered_text(doc, "信息工程学院本科生学年论文", 20, True, "黑体", before=12)
    add_centered_text(doc, MAIN_TITLE, 22, True, "黑体", before=70, after=8)
    add_centered_text(doc, SUB_TITLE, 16, False, "宋体", after=45)

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = ["姓　　名", "班　　级", "专　　业", "学　　院", "指导老师", "完成时间"]
    values = ["【姓名】", "【班级】", "【专业】", "信息工程学院", "【指导老师（姓名、职称）】", "【年　月　日】"]
    for row, label, value in zip(table.rows, labels, values):
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(9)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 100, 100, 100, 100)
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = borders.find(qn(f"w:{side}"))
                if tag is None:
                    tag = OxmlElement(f"w:{side}")
                    borders.append(tag)
                tag.set(qn("w:val"), "nil")
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p0.paragraph_format.first_line_indent = Cm(0)
        r0 = p0.add_run(label + "：")
        set_run_font(r0, size=Pt(14))
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.first_line_indent = Cm(0)
        r1 = p1.add_run(value)
        set_run_font(r1, size=Pt(14))

    doc.add_page_break()
    add_centered_text(doc, "学年论文评定表", 18, True, "黑体", before=4, after=14)
    eval_table = doc.add_table(rows=7, cols=2)
    eval_table.style = "Table Grid"
    eval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    eval_rows = [
        ("论文题目", MAIN_TITLE + SUB_TITLE),
        ("学生姓名", "【姓名】"),
        ("联系方式", "【联系方式】"),
        ("指导老师", "【指导老师（姓名、职称）】"),
        ("论文选题及成绩", "【由指导老师填写】"),
        ("指导老师评语", "\n\n\n\n\n【由指导老师填写】"),
        ("签名及日期", "指导老师签名：　　　　　　日期：　　年　月　日"),
    ]
    for idx, (label, value) in enumerate(eval_rows):
        eval_table.rows[idx].cells[0].width = Cm(4)
        eval_table.rows[idx].cells[1].width = Cm(11)
        for cell in eval_table.rows[idx].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 120, 120, 120, 120)
        p0 = eval_table.rows[idx].cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.first_line_indent = Cm(0)
        set_run_font(p0.add_run(label), size=Pt(12), bold=True)
        p1 = eval_table.rows[idx].cells[1].paragraphs[0]
        p1.paragraph_format.first_line_indent = Cm(0)
        set_run_font(p1.add_run(value), size=Pt(12))


def section_text(lines, start_marker, end_marker):
    start = lines.index(start_marker) + 1
    end = lines.index(end_marker)
    return [line for line in lines[start:end] if line.strip() not in {"---"}]


def add_abstracts_and_toc(doc: Document, lines: list[str]):
    doc.add_page_break()
    add_centered_text(doc, "摘　要", 18, True, "宋体", before=4, after=10)
    cn_lines = section_text(lines, "## 摘要", "## Abstract")
    cn_text = []
    keywords = ""
    for line in cn_lines:
        if line.startswith("**关键词"):
            keywords = "关键词：水质预测；动态图卷积；时间卷积网络；交叉注意力；汉江"
        elif line.strip():
            cn_text.append(clean_markdown(line))
    for text in cn_text:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0.85)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("关键词：")
    set_run_font(r, size=Pt(18), bold=True)
    r = p.add_run(keywords.removeprefix("关键词："))
    set_run_font(r, size=Pt(12))

    doc.add_page_break()
    add_centered_text(doc, "Abstract", 18, True, "Times New Roman", before=4, after=10)
    en_lines = section_text(lines, "## Abstract", "## 目录")
    en_text = []
    keywords_en = ""
    for line in en_lines:
        if line.startswith("**Key Words"):
            keywords_en = (
                "Keywords: water quality prediction; dynamic graph convolution; "
                "temporal convolutional network; cross-attention; Hanjiang River"
            )
        elif line.strip():
            en_text.append(clean_markdown(line))
    for text in en_text:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.85)
        r = p.add_run(text)
        set_run_font(r, western="Times New Roman", chinese="宋体", size=Pt(12))
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("Keywords: ")
    set_run_font(r, western="Times New Roman", chinese="宋体", size=Pt(15), bold=True)
    r = p.add_run(keywords_en.removeprefix("Keywords: "))
    set_run_font(r, western="Times New Roman", chinese="宋体", size=Pt(12))

    doc.add_page_break()
    add_centered_text(doc, "目　录", 18, True, "宋体", before=4, after=12)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    add_field(p, 'TOC \\o "1-2" \\h \\z \\u')
    note = doc.add_paragraph("提示：在 Word/WPS 中按 Ctrl+A 后按 F9 可更新目录与页码。")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.first_line_indent = Cm(0)
    note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    note.runs[0].font.size = Pt(9)


def locate_image(markdown_path: str) -> Path | None:
    override = SKILL_FIGURE_OVERRIDES.get(Path(markdown_path).name)
    if override is not None:
        return override if override.exists() else None
    candidate = ROOT / markdown_path
    if candidate.exists():
        return candidate
    basename = Path(markdown_path).name
    matches = list(FIGURE_ROOT.rglob(basename))
    return matches[0] if matches else None


def add_picture(doc: Document, path: Path, alt_text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.5))
    set_keep_with_next(p)


def add_markdown_table(doc: Document, rows: list[list[str]]):
    if len(rows) < 2:
        return
    data = [row for idx, row in enumerate(rows) if idx != 1]
    cols = max(len(row) for row in data)
    table = doc.add_table(rows=len(data), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(data):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            text = clean_markdown(row[c_idx]) if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=Pt(9.5), bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_equation(doc: Document, equation_lines: list[str]):
    raw = " ".join(part.strip() for part in equation_lines)
    tag_match = re.search(r"\\tag\{([^}]+)\}", raw)
    number = ""
    if tag_match:
        tag = tag_match.group(1)
        number = f"（式{EQUATION_MAP.get(tag, tag)}）"
        raw = raw[: tag_match.start()] + raw[tag_match.end() :]
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_math(p, raw.strip())
    if number:
        r2 = p.add_run("　" + number)
        set_run_font(r2, size=Pt(11))


def should_omit(line: str) -> bool:
    return any(token in line for token in OMIT_CONTAINS)


def add_body(doc: Document, lines: list[str]):
    start = lines.index("## 第1章 绪论")
    ref_start = lines.index("## 参考文献")
    body_lines = lines[start:ref_start]

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(section)
    set_page_number_start(section, 1)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.first_line_indent = Cm(0)
    hr = header.add_run("中南财经政法大学本科生学年论文")
    set_run_font(hr, size=Pt(9))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Cm(0)
    footer.add_run("-")
    add_field(footer, "PAGE")
    footer.add_run("-")

    i = 0
    in_mermaid = False
    skip_next_figure_caption = False
    while i < len(body_lines):
        line = body_lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_mermaid = not in_mermaid
            i += 1
            continue
        if in_mermaid or not stripped or stripped == "---" or should_omit(stripped):
            i += 1
            continue
        if "图5-1_多模型训练损失_同款.png" in stripped:
            i += 1
            continue

        if stripped.startswith("## "):
            title = academic_heading(stripped[3:], 1)
            doc.add_heading(title, level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            title = academic_heading(stripped[4:], 2)
            doc.add_heading(title, level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            title = academic_heading(stripped[5:], 3)
            doc.add_heading(title, level=3)
            i += 1
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            image_path = locate_image(image_match.group(2))
            if image_path:
                add_picture(doc, image_path, image_match.group(1))
                caption = doc.add_paragraph(style="Figure Caption")
                caption.add_run(clean_markdown(image_match.group(1)))
                set_keep_with_next(caption, False)
                skip_next_figure_caption = True
            i += 1
            continue

        if stripped == "$$":
            equation = []
            i += 1
            while i < len(body_lines) and body_lines[i].strip() != "$$":
                equation.append(body_lines[i])
                i += 1
            add_equation(doc, equation)
            i += 1
            continue

        if stripped.startswith("|"):
            rows = []
            while i < len(body_lines) and body_lines[i].strip().startswith("|"):
                rows.append(parse_table_row(body_lines[i]))
                i += 1
            add_markdown_table(doc, rows)
            continue

        text = clean_markdown(stripped)
        if not text:
            i += 1
            continue
        is_caption = stripped.startswith("**图") or stripped.startswith("**表")
        if is_caption:
            if text.startswith("图") and skip_next_figure_caption:
                skip_next_figure_caption = False
                i += 1
                continue
            p = doc.add_paragraph(style="Figure Caption")
            p.add_run(text)
            set_keep_with_next(p, text.startswith("表"))
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run(clean_markdown(stripped[2:]))
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run(clean_markdown(re.sub(r"^\d+\.\s*", "", stripped)))
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.widow_control = True
        i += 1

    reference_title = doc.add_heading("主要参考文献", level=1)
    for run in reference_title.runs:
        set_run_font(run, chinese="宋体", size=Pt(18), bold=True)
    references = []
    for line in lines[ref_start + 1 :]:
        stripped = line.strip()
        if re.match(r"^\[\d+\]", stripped):
            entry = re.sub(r"^\[\d+\]\s*", "", clean_markdown(stripped))
            references.append(entry)
    chinese = [entry for entry in references if re.search(r"[\u4e00-\u9fff]", entry.split(".")[0])]
    foreign = [entry for entry in references if entry not in chinese]
    for idx, entry in enumerate(chinese + foreign, 1):
        p = doc.add_paragraph(style="Reference")
        p.add_run(f"[{idx}] {entry}")


def finalize_document(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    core = doc.core_properties
    core.title = MAIN_TITLE + SUB_TITLE
    core.subject = "本科生学年论文"
    core.author = "【姓名】"
    core.keywords = "水质预测；动态图卷积；时间卷积网络；交叉注意力；汉江"
    core.comments = "由项目可核验数据与实验结果整理生成；封面信息待填写。"


def force_black_font_color(doc: Document):
    black = RGBColor(0, 0, 0)
    for style in doc.styles:
        if hasattr(style, "font"):
            style.font.color.rgb = black

    roots = [doc._element, doc.styles._element]
    for section in doc.sections:
        roots.extend(
            [
                section.header._element,
                section.first_page_header._element,
                section.even_page_header._element,
                section.footer._element,
                section.first_page_footer._element,
                section.even_page_footer._element,
            ]
        )

    for root in roots:
        for rpr in root.xpath(".//w:rPr"):
            color = rpr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                rpr.append(color)
            color.set(qn("w:val"), "000000")
            for attribute in ("themeColor", "themeTint", "themeShade"):
                color.attrib.pop(qn(f"w:{attribute}"), None)


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc.sections[0])
    doc.sections[0].header.is_linked_to_previous = False
    doc.sections[0].footer.is_linked_to_previous = False
    add_cover(doc)
    add_abstracts_and_toc(doc, lines)
    add_body(doc, lines)
    finalize_document(doc)
    force_black_font_color(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
